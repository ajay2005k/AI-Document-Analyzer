import os
from typing import List, Dict, Optional
from pathlib import Path
import PyPDF2
import docx
import numpy as np
from sentence_transformers import SentenceTransformer
from sklearn.metrics.pairwise import cosine_similarity
from transformers import T5Tokenizer, T5ForConditionalGeneration
from dataclasses import dataclass

@dataclass
class Document:
    content: str
    metadata: Dict
    embeddings: Optional[np.ndarray] = None

class DocumentQA:
    def __init__(self):
        # Initialize sentence transformer for embeddings (free to use)
        self.embedding_model = SentenceTransformer('all-MiniLM-L6-v2')
        # Initialize T5 model and tokenizer (free to use)
        self.tokenizer = T5Tokenizer.from_pretrained('t5-small')
        self.model = T5ForConditionalGeneration.from_pretrained('t5-small')
        self.documents: List[Document] = []
        self.chunk_size = 500
        self.chunk_overlap = 50
        self.max_context_length = 512  # T5 context length limit

    def load_document(self, file_path: str) -> None:
        """Load and process a document file."""
        content = self._read_file(file_path)
        chunks = self._chunk_text(content)
        
        for i, chunk in enumerate(chunks):
            doc = Document(
                content=chunk,
                metadata={'source': file_path, 'chunk_id': i}
            )
            # Generate embeddings for the chunk
            doc.embeddings = self.embedding_model.encode([chunk])[0]
            self.documents.append(doc)

    def _read_file(self, file_path: str) -> str:
        """Read different file formats and extract text content."""
        file_ext = Path(file_path).suffix.lower()
        
        if file_ext == '.pdf':
            return self._read_pdf(file_path)
        elif file_ext in ['.docx', '.doc']:
            return self._read_docx(file_path)
        elif file_ext in ['.txt', '.md']:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")

    def _read_pdf(self, file_path: str) -> str:
        """Extract text from PDF files."""
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            text = ''
            for page in reader.pages:
                text += page.extract_text() + '\n'
        return text

    def _read_docx(self, file_path: str) -> str:
        """Extract text from DOCX files."""
        doc = docx.Document(file_path)
        text = ''
        for paragraph in doc.paragraphs:
            text += paragraph.text + '\n'
        return text

    def _chunk_text(self, text: str) -> List[str]:
        """Split text into overlapping chunks."""
        words = text.split()
        chunks = []
        i = 0
        while i < len(words):
            chunk = ' '.join(words[i:i + self.chunk_size])
            chunks.append(chunk)
            i += self.chunk_size - self.chunk_overlap
        return chunks

    def _get_relevant_chunks(self, query: str, n_chunks: int = 3) -> List[Document]:
        """Retrieve the most relevant document chunks for a query."""
        query_embedding = self.embedding_model.encode([query])[0]
        
        similarities = []
        for doc in self.documents:
            similarity = cosine_similarity(
                [query_embedding],
                [doc.embeddings]
            )[0][0]
            similarities.append((similarity, doc))
        
        # Sort by similarity and get top n chunks
        similarities.sort(key=lambda x: x[0], reverse=True)
        return [doc for _, doc in similarities[:n_chunks]]

    def _construct_prompt(self, query: str, relevant_docs: List[Document]) -> str:
        """Construct a prompt for T5 using relevant chunks."""
        context = " ".join([doc.content for doc in relevant_docs])
        # Format for T5: "question: {question} context: {context}"
        return f"question: {query} context: {context}"

    def answer_question(self, query: str) -> Dict:
        """Generate an answer for a question using the loaded documents."""
        # Get relevant document chunks
        relevant_chunks = self._get_relevant_chunks(query)
        
        # Construct the prompt
        prompt = self._construct_prompt(query, relevant_chunks)
        
        # Tokenize and generate answer using T5
        inputs = self.tokenizer.encode(prompt, return_tensors="pt", max_length=self.max_context_length, truncation=True)
        outputs = self.model.generate(inputs, max_length=150, min_length=40, length_penalty=2.0, num_beams=4)
        answer = self.tokenizer.decode(outputs[0], skip_special_tokens=True)
        
        # Return answer with metadata
        return {
            'answer': answer,
            'sources': [doc.metadata['source'] for doc in relevant_chunks],
            'confidence': len(relevant_chunks) > 0  # Simple confidence measure
        }

def interactive_mode(qa_system):
    print("\nEntering interactive mode. Type 'quit' to exit.")
    while True:
        question = input("\nEnter your question: ")
        if question.lower() == 'quit':
            break
            
        result = qa_system.answer_question(question)
        print("\nAnswer:", result['answer'])
        print("Sources:", result['sources'])
        print("Confidence:", result['confidence'])

def select_file():
    """Let user select the file type and path."""
    print("\nSelect file type:")
    print("1. Text file (.txt)")
    print("2. PDF file (.pdf)")
    
    while True:
        choice = input("\nEnter your choice (1 or 2): ")
        if choice in ['1', '2']:
            break
        print("Invalid choice. Please enter 1 or 2.")
    
    if choice == '1':
        file_extension = '.txt'
        print("\nMake sure your text file is in the current directory.")
    else:
        file_extension = '.pdf'
        print("\nMake sure your PDF file is in the current directory.")
    
    # List available files
    current_dir = os.getcwd()
    available_files = [f for f in os.listdir(current_dir) if f.endswith(file_extension)]
    
    if not available_files:
        print(f"\nNo {file_extension} files found in the current directory!")
        print(f"Current directory: {current_dir}")
        return None
    
    print("\nAvailable files:")
    for i, file in enumerate(available_files, 1):
        print(f"{i}. {file}")
    
    while True:
        try:
            file_num = int(input("\nEnter the number of the file you want to analyze: "))
            if 1 <= file_num <= len(available_files):
                return available_files[file_num - 1]
            print("Invalid number. Please try again.")
        except ValueError:
            print("Please enter a valid number.")

def main():
    qa_system = DocumentQA()
    
    # Let user select the file
    selected_file = select_file()
    if not selected_file:
        print("No file selected. Exiting...")
        return
    
    print(f"\nLoading document: {selected_file}")
    try:
        qa_system.load_document(selected_file)
        print("Document loaded successfully!")
        
        # Enter interactive mode
        interactive_mode(qa_system)
    except Exception as e:
        print(f"Error loading document: {e}")

if __name__ == "__main__":
    main()
